import tkinter as tk
from tkinter import filedialog
from tkinter import messagebox
import os
import csv
from docx import Document
import re

def extract_urls_from_docx(docx_file):
    document = Document(docx_file)
    urls = []

    # Extract URLs from paragraphs
    for paragraph in document.paragraphs:
        urls.extend(re.findall(r'https?://\S+', paragraph.text))

    # Extract URLs from tables
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                urls.extend(re.findall(r'https?://\S+', cell.text))

    return urls



def convert_to_csv(docx_file, csv_file):
    urls = extract_urls_from_docx(docx_file)
    with open(csv_file, 'w', newline='') as csvfile:
        url_writer = csv.writer(csvfile)
        url_writer.writerow(['URLs'])
        for url in urls:
            url_writer.writerow([url])

def browse_file():
    file_path = filedialog.askopenfilename(filetypes=[("Word files", "*.docx")])
    if file_path:
        entry.delete(0, tk.END)
        entry.insert(tk.END, file_path)

def convert_and_save():
    docx_file = entry.get()
    if not docx_file:
        messagebox.showerror("Error", "Please select a .docx file.")
        return

    csv_file = os.path.splitext(docx_file)[0] + ".csv"
    convert_to_csv(docx_file, csv_file)
    messagebox.showinfo("Success", f"CSV file saved at {csv_file}")

# GUI
root = tk.Tk()
root.title("SHOLEF - DOCX links extractor")

frame = tk.Frame(root)
frame.pack(padx=10, pady=10)

label = tk.Label(frame, text="Select a .docx file:")
label.grid(row=0, column=0, sticky="w")

entry = tk.Entry(frame, width=50)
entry.grid(row=0, column=1)

browse_button = tk.Button(frame, text="Browse", command=browse_file)
browse_button.grid(row=0, column=2)

convert_button = tk.Button(frame, text="Convert to CSV", command=convert_and_save)
convert_button.grid(row=1, column=0, columnspan=3, pady=10)

root.mainloop()
